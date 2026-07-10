from docx import Document

summary = Document()
summary.add_heading('Final Changes Summary', level=1)
summary.add_paragraph('This document captures the important final changes made since this afternoon. Each change is shown with a before/after description.')

changes = [
    {
        'title': 'Navigation Order',
        'before': 'About Us appeared before RCN Retreat in the site header navigation.',
        'after': 'About Us now appears between RCN Retreat and Contact in the site header navigation.',
    },
    {
        'title': 'Hero Section',
        'before': 'HEALING WITH PURPOSE was shown as a normal text line below the logo.',
        'after': 'HEALING WITH PURPOSE is now bold, highlighted, and placed directly below the logo in the hero section.',
    },
    {
        'title': 'Therapy Catalogue Layout',
        'before': 'Therapy images were large and always visible in a full-size carousel layout for each therapy.',
        'after': 'Therapy names render as list-style rows; hovering a row expands the therapy image carousel, and clicking View Details expands the full detail panel.',
    },
    {
        'title': 'Therapy View Details QR',
        'before': 'QR booking codes were not consistently shown under every therapy detail section.',
        'after': 'A booking QR code now appears below each therapy View Details panel for easier booking.',
    },
    {
        'title': 'Info Icon Behavior',
        'before': 'Therapy detail info text was always visible under the info icon.',
        'after': 'Clicking the info icon now toggles the descriptive text for each therapy detail item.',
    },
    {
        'title': 'Therapy Model',
        'before': 'Therapy data included only a single image and an optional hasQR flag.',
        'after': 'Therapy data now supports an images array for carousel slides and an optional qrValue field for booking links.',
    },
]

for change in changes:
    summary.add_heading(change['title'], level=2)
    summary.add_paragraph('Before:', style='List Bullet')
    summary.add_paragraph(change['before'], style='BodyText')
    summary.add_paragraph('After:', style='List Bullet')
    summary.add_paragraph(change['after'], style='BodyText')

summary.save('d:/DRT/final-changes-summary.docx')
print('Created d:/DRT/final-changes-summary.docx')
